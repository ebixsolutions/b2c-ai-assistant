# -*- coding: utf-8 -*-
"""
Generate `AI_Product_Recommendation_Rules.docx` — full-detail documentation for
the 9 deterministic recommendation rules (1,2,3,4,5,7,8,9,10; Rule 6 skipped:
it needs an external Google Trends signal that is not configured).

Content is transcribed verbatim from the verified source of truth:
  - app/rules/queries.py   (the actual SQL formulas)
  - app/rules/registry.py  (tags, labels, KPI targets, reasons)
  - docs/rule_*.md         (worked examples, output fields, tuning knobs)

Run:  python generate_rules_docx.py
"""
from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT

ACCENT = RGBColor(0x1F, 0x4E, 0x79)   # dark blue
MONO_BG = "F2F2F2"

doc = Document()

# ---- base styles -----------------------------------------------------------
normal = doc.styles["Normal"]
normal.font.name = "Calibri"
normal.font.size = Pt(11)


# ---- helpers ---------------------------------------------------------------
def h1(text):
    p = doc.add_heading(text, level=1)
    for r in p.runs:
        r.font.color.rgb = ACCENT
    return p


def h2(text):
    p = doc.add_heading(text, level=2)
    for r in p.runs:
        r.font.color.rgb = ACCENT
    return p


def h3(text):
    p = doc.add_heading(text, level=3)
    for r in p.runs:
        r.font.color.rgb = ACCENT
    return p


def para(text="", bold=False, italic=False):
    p = doc.add_paragraph()
    r = p.add_run(text)
    r.bold = bold
    r.italic = italic
    return p


def bullet(text):
    return doc.add_paragraph(text, style="List Bullet")


def code_block(text):
    """Monospaced, shaded single-cell table — renders like a code box in Word."""
    tbl = doc.add_table(rows=1, cols=1)
    tbl.alignment = WD_TABLE_ALIGNMENT.LEFT
    cell = tbl.cell(0, 0)
    # shading
    from docx.oxml.ns import qn
    from docx.oxml import OxmlElement
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:fill"), MONO_BG)
    cell._tc.get_or_add_tcPr().append(shd)
    cell.text = ""
    first = True
    for line in text.strip("\n").split("\n"):
        p = cell.paragraphs[0] if first else cell.add_paragraph()
        first = False
        r = p.add_run(line)
        r.font.name = "Consolas"
        r.font.size = Pt(9.5)
        p.paragraph_format.space_after = Pt(0)
        p.paragraph_format.space_before = Pt(0)
    doc.add_paragraph()
    return tbl


def insight_box(text):
    """Soft-pink shaded single-cell box that mimics the live 'AI INSIGHT' UI panel."""
    from docx.oxml.ns import qn
    from docx.oxml import OxmlElement
    tbl = doc.add_table(rows=1, cols=1)
    tbl.alignment = WD_TABLE_ALIGNMENT.LEFT
    cell = tbl.cell(0, 0)
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:fill"), "FCEEF0")  # soft pink, like the screenshot
    cell._tc.get_or_add_tcPr().append(shd)
    cell.text = ""
    p = cell.paragraphs[0]
    r = p.add_run(text)
    r.font.size = Pt(10.5)
    p.paragraph_format.space_after = Pt(0)
    p.paragraph_format.space_before = Pt(0)
    doc.add_paragraph()
    return tbl


def make_table(headers, rows):
    tbl = doc.add_table(rows=1, cols=len(headers))
    tbl.style = "Light Grid Accent 1"
    tbl.alignment = WD_TABLE_ALIGNMENT.LEFT
    hdr = tbl.rows[0].cells
    for i, htext in enumerate(headers):
        hdr[i].text = ""
        r = hdr[i].paragraphs[0].add_run(htext)
        r.bold = True
        r.font.size = Pt(10)
    for row in rows:
        cells = tbl.add_row().cells
        for i, val in enumerate(row):
            cells[i].text = ""
            r = cells[i].paragraphs[0].add_run(str(val))
            r.font.size = Pt(10)
    doc.add_paragraph()
    return tbl


# ============================================================================
# COVER / TITLE
# ============================================================================
t = doc.add_paragraph()
t.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = t.add_run("AI Product Recommendation Engine")
run.bold = True
run.font.size = Pt(26)
run.font.color.rgb = ACCENT

st = doc.add_paragraph()
st.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = st.add_run("Business Rule Documentation — Rules 1, 2, 3, 4, 5, 7, 8, 9, 10")
r.font.size = Pt(13)
r.italic = True

doc.add_paragraph()

# ============================================================================
# OVERVIEW
# ============================================================================
h1("Overview")
para(
    "We have approximately 100,000 products in the database. Sending all product "
    "data directly to an AI model is not practical because it exceeds token limits, "
    "increases cost, slows response times, and can lead to hallucinations or "
    "inconsistent recommendations."
)
para(
    "To address this, each recommendation category is defined using a deterministic "
    "business formula provided by the business team. The database performs the heavy "
    "filtering and ranking, while AI focuses only on evaluating a small set of top "
    "products and generating explainable recommendations."
)

# ---- process flow ----------------------------------------------------------
h1("Process Flow")
code_block(
    "100,000 Products\n"
    "        |\n"
    "Rule-Based Filtering\n"
    "        |\n"
    "~500 Eligible Products\n"
    "        |\n"
    "SQL Ranking Formula\n"
    "        |\n"
    "Top 2-10 Products\n"
    "        |\n"
    "AI (GPT / Gemini) Evaluation\n"
    "        |\n"
    "1 Recommended Product"
)

# ---- key principle ---------------------------------------------------------
h1("Key Principle")
para(
    "Deterministic database logic performs the heavy filtering and ranking, while "
    "AI focuses only on reasoning over a small set of high-quality products. "
    "SQL selects the products; AI does NOT pick, rank, or drop anything — it only "
    "rewrites the detection logic and recommendation reason using the store's real "
    "numbers. If AI is disabled or fails, the static specification text is returned "
    "verbatim."
)
para("Database Responsibilities:", bold=True)
for x in ["Apply business rules", "Filter products", "Calculate ranking scores",
          "Select the top N products (default 2, configurable via PRODUCT_LIMIT)"]:
    bullet(x)
para("AI Responsibilities:", bold=True)
for x in ["Compare shortlisted products", "Generate explanations",
          "Recommend the best product", "Produce human-readable insights"]:
    bullet(x)

doc.add_page_break()

# ============================================================================
# RULE DATA (verbatim from queries.py / registry.py / docs)
# ============================================================================
RULES = [
    {
        "id": 1,
        "title": "Potential Viral Product",
        "tag": "viral",
        "kpi": "Conversion Rate (CR) +15%",
        "detects": (
            "A product that is being discovered through external traffic (ads, social, "
            "referrals) and shows unusually strong engagement (high cart-add rate) despite "
            "low page-view traffic. These are early-signal viral candidates worth amplifying."
        ),
        "formula": (
            "CTR Proxy  >  Store Avg CTR x 1.5\n"
            "AND\n"
            "PV         <  Store Avg PV  x 0.5\n"
            "AND\n"
            "type       =  16            (external click in user_tag_define)"
        ),
        "steps": [
            ("Step 1 - Find externally clicked products",
             "external_products = user_tag_define\n"
             "                    WHERE type = 16\n"
             "                      AND created_at >= NOW - window_days"),
            ("Step 2 - Map pixel events to local items",
             "Join shopify_web_pixel_log to item_master_attribute via variantId\n"
             "-> resolves each pixel row to a local item_id."),
            ("Step 3 - Keep only externally discovered items",
             "Intersect Step 1 + Step 2 - only items that were both externally\n"
             "clicked and saw pixel activity."),
            ("Step 4 - Per-item counts",
             "pv  = COUNT(event = product_viewed)\n"
             "atc = COUNT(event = product_added_to_cart)"),
            ("Step 5 - CTR Proxy (per item)",
             "ctr_proxy = atc / pv          (0 when pv = 0)"),
            ("Step 6 - Store averages (across all candidates)",
             "avg_ctr = AVG(ctr_proxy) OVER ()\n"
             "avg_pv  = AVG(pv)        OVER ()"),
            ("Step 7 - Filter (the two rule conditions)",
             "ctr_proxy >  1.5 x avg_ctr         (high engagement)\n"
             "AND\n"
             "pv        <  0.5 x avg_pv          (low traffic)"),
            ("Step 8 - Rank",
             "ctr_ratio = ctr_proxy / avg_ctr    (how many x above average CTR)\n"
             "Order by ctr_ratio DESC, return top PRODUCT_LIMIT."),
        ],
        "example_note": "Suppose avg_ctr = 0.20 and avg_pv = 40. Threshold A = 1.5 x 0.20 = 0.30. Threshold B = 0.5 x 40 = 20.",
        "example_headers": ["item", "pv", "atc", "ctr_proxy = atc/pv", "ctr_proxy > 0.30?", "pv < 20?", "passes?"],
        "example_rows": [
            ["A", "10", "5", "5/10 = 0.500", "Yes", "Yes", "QUALIFIES"],
            ["B", "80", "30", "30/80 = 0.375", "Yes", "No", "no"],
            ["C", "5", "0", "0/5 = 0.000", "No", "Yes", "no"],
        ],
        "example_concl": "Item A wins: low traffic but converts at 0.50 / 0.20 = 2.5x the store average.",
        "outputs": [
            ("page_views", "pv count in window"),
            ("add_to_carts", "atc count in window"),
            ("ctr_proxy", "atc / pv"),
            ("store_avg_ctr", "mean ctr_proxy across candidates"),
            ("score", "ctr_proxy / avg_ctr (how many x above average)"),
            ("traffic_window_days", "rolling window used (default 30)"),
        ],
    },
    {
        "id": 2,
        "title": "Inventory Clearance Candidate",
        "tag": "clearance",
        "kpi": "Inventory Turnover +25%",
        "detects": (
            "A product with severe inventory accumulation: a large quantity on hand but "
            "very little recent movement. These items tie up capital and warehouse space "
            "and are prime candidates for promotional campaigns to liquidate."
        ),
        "formula": (
            "total_inventory  >  100 units\n"
            "AND\n"
            "( sales_30d < 5  OR  days_of_inventory > 90 )"
        ),
        "steps": [
            ("Step 1 - Inventory on hand (stock CTE)",
             "total_qty = SUM(quantity) FROM stock_inventroy  GROUP BY item_id"),
            ("Step 2 - 30-day sales (sales30 CTE)",
             "qty_30d = SUM(quantity) FROM sales_order_details d\n"
             "          JOIN sales_order o ON o.id = d.sales_id\n"
             "          WHERE o.created_at >= NOW - 30 days"),
            ("Step 3 - Days of inventory (days of cover)",
             "doi_days = total_qty / (qty_30d / 30.0)     (NULL when qty_30d = 0)"),
            ("Step 4 - Filter (the rule conditions)",
             "total_qty > 100\n"
             "AND ( qty_30d < 5  OR  COALESCE(doi_days, 9999) > 90 )"),
            ("Step 5 - Rank",
             "Order by total_qty DESC, return top PRODUCT_LIMIT."),
        ],
        "example_note": "Threshold A = inventory > 100. Threshold B = 30-day sales < 5 OR days of inventory > 90.",
        "example_headers": ["item", "total_qty", "qty_30d", "doi_days", "qty>100?", "<5 or doi>90?", "passes?"],
        "example_rows": [
            ["A", "500", "3", "5000", "Yes", "Yes (3 < 5)", "QUALIFIES"],
            ["B", "800", "20", "1200", "Yes", "Yes (doi > 90)", "QUALIFIES"],
            ["C", "60", "1", "1800", "No", "Yes", "no (inventory <= 100)"],
        ],
        "example_concl": "Item A and B both carry heavy stock that is barely moving - liquidate first by largest quantity.",
        "outputs": [
            ("inventory", "total quantity on hand (SUM of stock_inventroy.quantity)"),
            ("sales_30d", "units sold in the last 30 days"),
            ("days_of_inventory", "total_qty / (qty_30d / 30) - days of cover"),
            ("score", "the inventory quantity (used for ranking)"),
        ],
        "note": (
            "Data dependency: requires stock_inventroy to be populated (the table is often "
            "empty -> rule returns 0). Note the table name is spelled 'stock_inventroy'. The "
            "100-unit / 5-sales / 90-day thresholds are hard-coded in queries.py."
        ),
    },
    {
        "id": 3,
        "title": "Hidden Gem Product",
        "tag": "hidden_gem",
        "kpi": "ROAS +30%",
        "detects": (
            "A product that converts at an extremely high rate (much better than average) "
            "but almost nobody sees it. These products are under-exposed wins - give them "
            "more impressions/spend and ROAS jumps."
        ),
        "formula": (
            "CR             >  Store Avg CR x 2.0\n"
            "AND\n"
            "impressions    in bottom 20% by traffic"
        ),
        "steps": [
            ("Step 1 - Impressions per item",
             "impressions = COUNT(event = product_viewed)"),
            ("Step 2 - Orders per item",
             "orders_30d = COUNT(DISTINCT sales_id)"),
            ("Step 3 - Conversion rate (the key division)",
             "cr = orders_30d / impressions"),
            ("Step 4 - Impression percentile (low-traffic detector)",
             "impressions_pct = PERCENT_RANK() OVER (ORDER BY impressions ASC)\n"
             "A value of 0.20 means the item has fewer impressions than 80% of items."),
            ("Step 5 - Store-wide average CR",
             "avg_cr = AVG(cr) OVER ()"),
            ("Step 6 - Filter (the two rule conditions)",
             "cr               >  2.0 x avg_cr     (converts at 2x store average)\n"
             "AND\n"
             "impressions_pct  <= 0.20             (bottom 20% of traffic)"),
            ("Step 7 - Rank",
             "Order by cr DESC, return top PRODUCT_LIMIT."),
        ],
        "example_note": "Suppose store avg_cr = 0.04 (4%) and there are 10 items. Threshold A = 2.0 x 0.04 = 0.08. Threshold B = impressions_pct <= 0.20.",
        "example_headers": ["item", "impressions", "orders_30d", "cr = orders/impr", "impr_pct", "cr>0.08?", "pct<=0.20?", "passes?"],
        "example_rows": [
            ["A", "12", "2", "2/12 = 0.167", "0.00", "Yes", "Yes", "QUALIFIES"],
            ["B", "18", "1", "1/18 = 0.056", "0.11", "No", "Yes", "no"],
            ["C", "200", "30", "30/200 = 0.150", "0.78", "Yes", "No", "no"],
        ],
        "example_concl": "Item A: 1 buyer in every 6 viewers (0.167 / 0.04 = 4.2x store average), only 12 impressions. Classic hidden gem.",
        "outputs": [
            ("impressions", "product_viewed count (30d)"),
            ("orders_30d", "distinct orders containing the item (30d)"),
            ("conversion_rate", "orders_30d / impressions"),
            ("store_avg_cr", "mean CR across candidate items"),
            ("score", "the cr value (used for ranking)"),
        ],
    },
    {
        "id": 4,
        "title": "High Attention, Low Purchase",
        "tag": "high_attention",
        "kpi": "Add-to-Cart Rate (ATC) +20%",
        "detects": (
            "A product that many shoppers view but few add to cart - the page is pulling "
            "traffic but something is killing conversion (price, photos, description, "
            "reviews, stock). Classic candidate for a discount, page refresh, or A/B test."
        ),
        "formula": (
            "PV in Top 5% by traffic\n"
            "AND\n"
            "ATC rate < Store Avg ATC rate x 0.6"
        ),
        "steps": [
            ("Step 1 - Pixel -> item join",
             "Pull all shopify_web_pixel_log rows, resolve variantId -> local item_id\n"
             "via item_master_attribute."),
            ("Step 2 - Per-item counts",
             "pv  = COUNT(event = product_viewed)\n"
             "atc = COUNT(event = product_added_to_cart)"),
            ("Step 3 - Per-item ATC rate (key division)",
             "atc_rate = atc / pv          (0 when pv = 0)"),
            ("Step 4 - Store-wide stats",
             "n_items      = COUNT(*)\n"
             "avg_atc_rate = AVG(atc_rate)"),
            ("Step 5 - PV ranking with ties",
             "pv_rank = RANK() OVER (ORDER BY pv DESC)\n"
             "Items tied on pv share the same rank - critical for small datasets."),
            ("Step 6 - Top-5% cutoff (with safety floor)",
             "top_n_cutoff = GREATEST( PRODUCT_LIMIT, CEIL(n_items x 0.05) )"),
            ("Step 7 - Filter (the two rule conditions)",
             "pv_rank  <=  top_n_cutoff               (high traffic - top 5%)\n"
             "AND\n"
             "atc_rate <   0.6 x avg_atc_rate         (clearly below-average cart-add)"),
            ("Step 8 - Order and limit",
             "Order by pv DESC, then LIMIT PRODUCT_LIMIT."),
        ],
        "example_note": "n_items = 4, avg_atc_rate = 0.250. Threshold = 0.6 x 0.250 = 0.150. Cutoff = GREATEST(2, CEIL(4 x 0.05)=1) = 2.",
        "example_headers": ["item_id", "pv", "atc", "atc_rate", "rate<0.150?", "pv_rank", "rank<=2?", "passes?"],
        "example_rows": [
            ["1033458", "4", "4", "1.000", "No", "4", "No", "no"],
            ["1033468", "25", "0", "0.000", "Yes", "1", "Yes", "QUALIFIES"],
            ["1033460", "21", "0", "0.000", "Yes", "2", "Yes", "QUALIFIES"],
            ["1033471", "21", "0", "0.000", "Yes", "2", "Yes", "QUALIFIES"],
        ],
        "example_concl": "Both pv=21 rows tie at pv_rank = 2 (the point of RANK()). 3 rows pass the WHERE clause; final LIMIT 2 trims to 2 displayed.",
        "outputs": [
            ("page_views", "pv (30d)"),
            ("add_to_carts", "atc (30d)"),
            ("atc_rate", "atc / pv"),
            ("store_avg_atc", "avg_atc_rate across candidate items"),
            ("score", "the pv value (used for ordering)"),
        ],
        "note": (
            "RANK() is used (not NTILE/PERCENTILE_CONT) because it gives integer ranks and "
            "shares the rank on ties, which matches what humans expect from 'top N'."
        ),
    },
    {
        "id": 5,
        "title": "Shopping Basket Magnet",
        "tag": "basket_magnet",
        "kpi": "Average Order Value (AOV) +18%",
        "detects": (
            "A product that is rarely bought alone - when shoppers buy it, they tend to add "
            "other items to the same order. These are natural anchors for bundles, "
            "'frequently bought together' widgets, and cross-sell campaigns."
        ),
        "formula": (
            "multi_product_orders / total_orders  >  0.40\n"
            "(more than 40% of this item's orders contained at least one other distinct item)"
        ),
        "steps": [
            ("Step 1 - Order sizes (order_sizes CTE)",
             "n_items_per_order = COUNT(DISTINCT item_id)  GROUP BY sales_id"),
            ("Step 2 - Per-item order counts (per_item CTE)",
             "total_orders = COUNT(DISTINCT sales_id)\n"
             "multi_orders = COUNT(DISTINCT sales_id) WHERE n_items_per_order > 1"),
            ("Step 3 - Multi-product ratio (key division)",
             "multi_ratio = multi_orders / total_orders"),
            ("Step 4 - Filter (the rule condition)",
             "total_orders >  0\n"
             "AND\n"
             "multi_ratio  >  0.40"),
            ("Step 5 - Rank",
             "Order by multi_ratio DESC, then total_orders DESC (tiebreaker).\n"
             "Return top PRODUCT_LIMIT."),
        ],
        "example_note": "Threshold = multi_ratio > 0.40.",
        "example_headers": ["item", "total_orders", "multi_orders", "multi_ratio = multi/total", "> 0.40?", "passes?"],
        "example_rows": [
            ["A", "50", "40", "40/50 = 0.80", "Yes", "QUALIFIES"],
            ["B", "100", "35", "35/100 = 0.35", "No", "no"],
            ["C", "5", "5", "5/5 = 1.00", "Yes", "QUALIFIES (ranks below A by tie-break)"],
        ],
        "example_concl": "Item A wins: high volume and shows up in baskets with other products 80% of the time.",
        "outputs": [
            ("total_orders", "distinct orders containing the item"),
            ("multi_product_orders", "of those, how many had >=2 distinct items"),
            ("multi_order_ratio", "multi_orders / total_orders"),
            ("score", "the multi_ratio value (used for ranking)"),
        ],
        "note": "No time window - uses all-time order history. The 0.4 threshold is hard-coded in queries.py.",
    },
    {
        "id": 7,
        "title": "Customer Loyalty Favorite",
        "tag": "loyalty",
        "kpi": "LTV (Lifetime Value) +22%",
        "detects": (
            "A product that customers buy more than once. A high repeat-buyer rate means "
            "people came back specifically for this item - these are LTV anchors, perfect "
            "for subscription, replenishment, or VIP campaigns."
        ),
        "formula": (
            "repeat_buyers / total_buyers  >  0.20\n"
            "(more than 20% of buyers placed at least 2 orders containing this item)"
        ),
        "steps": [
            ("Step 1 - Per-(item, buyer) order counts (item_buyers CTE)",
             "Join sales_order_details -> sales_order -> user on create_user_id.\n"
             "Filter to real B2C shoppers: user.relation_type = 'b2c'\n"
             "AND user.relation_company_id = tenant.\n"
             "orders_by_buyer = COUNT(DISTINCT sales_id) GROUP BY item_id, user_id"),
            ("Step 2 - Per-item buyer aggregates (per_item CTE)",
             "total_buyers  = COUNT(*)                              (distinct buyers)\n"
             "repeat_buyers = COUNT(*) WHERE orders_by_buyer >= 2   (came back >=2 orders)"),
            ("Step 3 - Repeat rate (key division)",
             "repeat_rate = repeat_buyers / total_buyers"),
            ("Step 4 - Filter (the rule condition)",
             "total_buyers >  0\n"
             "AND\n"
             "repeat_rate  >  0.20"),
            ("Step 5 - Rank",
             "Order by repeat_rate DESC, then total_buyers DESC (tiebreaker).\n"
             "Return top PRODUCT_LIMIT."),
        ],
        "example_note": "Threshold = repeat_rate > 0.20.",
        "example_headers": ["item", "total_buyers", "repeat_buyers", "repeat_rate = repeat/total", "> 0.20?", "passes?"],
        "example_rows": [
            ["A", "100", "35", "35/100 = 0.35", "Yes", "QUALIFIES"],
            ["B", "200", "30", "30/200 = 0.15", "No", "no"],
            ["C", "4", "3", "3/4 = 0.75", "Yes", "QUALIFIES (tiny sample, low confidence)"],
        ],
        "example_concl": "Item A wins on volume + retention: 35 of 100 buyers came back for more.",
        "outputs": [
            ("total_customers", "distinct buyers ever"),
            ("repeat_customers", "of those, how many ordered >=2 times"),
            ("repeat_rate", "repeat_customers / total_customers"),
            ("score", "the repeat_rate value (used for ranking)"),
        ],
        "note": (
            "No time window - uses all-time purchase history. Customer identity comes "
            "directly from sales_order.create_user_id -> user.id. The 0.20 threshold is "
            "hard-coded in queries.py."
        ),
    },
    {
        "id": 8,
        "title": "Profit Protection Product",
        "tag": "profit",
        "kpi": "Net Profit +12%",
        "detects": (
            "A product with the highest profit potential: gross margin above 50% over the "
            "last 90 days. Suitable for higher advertising budgets because each sale "
            "contributes strong profit."
        ),
        "formula": (
            "gross_margin = (Selling Price - Cost) / Selling Price  >  50%\n"
            "AND  cost > 0   (over the last 90 days)"
        ),
        "steps": [
            ("Step 1 - Primary cost (inv_cost CTE)",
             "avg_inv_cost = AVG(cost) FROM stock_inventroy WHERE cost > 0\n"
             "(authoritative unit cost held in inventory - used first)"),
            ("Step 2 - Last-resort cost (purchase_cost CTE)",
             "avg_purchase_cost = AVG(unit_price) FROM item_master_purchase\n"
             "WHERE unit_price > 0"),
            ("Step 3 - Cost flow (first non-zero wins, 100% coverage)",
             "stock_inventroy.cost -> sales_order_details.cost\n"
             "-> item_master_purchase.unit_price -> 0"),
            ("Step 4 - Per-item margin (per_item CTE, last 90 days)",
             "avg_margin = AVG( (price - cost) / price )\n"
             "where cost = COALESCE(inv_cost, sales-line cost, purchase cost, 0)"),
            ("Step 5 - Filter (the rule conditions)",
             "avg_margin > 0.5\n"
             "AND\n"
             "avg_cost   > 0"),
            ("Step 6 - Rank",
             "Order by avg_margin DESC, return top PRODUCT_LIMIT."),
        ],
        "example_note": "Threshold = gross margin > 0.50 (50%).",
        "example_headers": ["item", "avg_price", "avg_cost", "margin = (price-cost)/price", "> 0.50?", "passes?"],
        "example_rows": [
            ["A", "100", "30", "(100-30)/100 = 0.70", "Yes", "QUALIFIES"],
            ["B", "80", "50", "(80-50)/80 = 0.375", "No", "no"],
            ["C", "120", "0", "n/a (cost = 0)", "-", "no (cost not known)"],
        ],
        "example_concl": "Item A keeps 70 cents of every sales dollar - safe to spend more on ads.",
        "outputs": [
            ("gross_margin", "(price - cost) / price"),
            ("gross_margin_pct", "gross_margin x 100"),
            ("avg_selling_price", "average selling price over 90 days"),
            ("avg_cost", "average resolved unit cost"),
            ("qty_sold_90d", "units sold in the last 90 days"),
            ("score", "the avg_margin value (used for ranking)"),
        ],
        "note": (
            "Cost resolution priority: stock_inventroy.cost (per-item unit cost) first, then "
            "the sales-line cost, then the supplier purchase price. If all three are "
            "0/missing, margin cannot be computed and the item is excluded."
        ),
    },
    {
        "id": 9,
        "title": "User Engagement Leader",
        "tag": "engagement",
        "kpi": "Dwell Time +35%",
        "detects": (
            "Products whose visitors spend a long time on the product page - average dwell "
            "time greater than 90 seconds per visitor. A high dwell time signals strong "
            "content attraction and genuine interest, independent of how many people viewed "
            "the page."
        ),
        "formula": "qualifies if  avg_dwell_seconds > 90",
        "steps": [
            ("Step 1 - View-end events (view_ends CTE)",
             "Pull all product_view_end rows with their reason + (for tab_closed)\n"
             "the startedAt / endedAt JS-ms timestamps from event_data."),
            ("Step 2 - Branch A: reason = 'page_exit' (row timestamps, epoch seconds)",
             "Pair the product_view_end with the immediately-preceding product_viewed\n"
             "row BY id (largest id below, same user+item):\n"
             "dwell = view_end.created_at - viewed.created_at"),
            ("Step 3 - Branch B: reason = 'tab_closed' (JS ms timestamps in event_data)",
             "dwell = (endedAt - startedAt) / 1000\n"
             "Rows missing startedAt or endedAt are dropped."),
            ("Step 4 - Per visitor",
             "user_dwell = AVG(dwell) GROUP BY (item_id, user_id)\n"
             "(a repeat viewer counts as a single visitor, not once per page view)"),
            ("Step 5 - Per item",
             "avg_dwell = AVG(user_dwell) GROUP BY item_id\n"
             "customers = COUNT(distinct visitors)"),
            ("Step 6 - Filter and rank",
             "WHERE avg_dwell > 90\n"
             "Order by avg_dwell DESC, return top PRODUCT_LIMIT."),
        ],
        "example_note": "Threshold = avg_dwell_seconds > 90.",
        "example_headers": ["item", "customers", "avg_dwell (s)", "> 90s?", "passes?"],
        "example_rows": [
            ["A", "12", "145", "Yes", "QUALIFIES"],
            ["B", "30", "62", "No", "no"],
            ["C", "5", "210", "Yes", "QUALIFIES (ranks first by dwell)"],
        ],
        "example_concl": "Items A and C hold visitors well past 90 seconds - strong content attraction worth surfacing.",
        "outputs": [
            ("avg_dwell_seconds", "average dwell time across distinct visitors (1 dp)"),
            ("customers", "number of distinct engaged visitors"),
            ("score", "the avg_dwell value (used for ranking)"),
        ],
        "note": (
            "Dwell time is NOT a stored column - it is derived from product_view_end pixel "
            "events. This rule requires the web pixel to emit product_view_end carrying the "
            "same event_data.variantId as the matching product_viewed; until then no session "
            "pairs and the rule returns 0. Rule 9 uses its own pixel CTE that carries "
            "event_data through (the shared CTE does not expose it)."
        ),
    },
    {
        "id": 10,
        "title": "New Product Momentum",
        "tag": "momentum",
        "kpi": "New Customer Acquisition +20%",
        "detects": (
            "A newly launched product (created in the last 7 days) whose today's sales pace "
            "is more than 2x its 7-day average. Signal: this product is accelerating - "
            "promote it before momentum dies."
        ),
        "formula": (
            "item created within last 7 days\n"
            "AND  qty_7d > 0\n"
            "AND  today_qty  >  2 x (qty_7d / 7)"
        ),
        "steps": [
            ("Step 1 - New items (new_items CTE)",
             "item_master.created_at >= now - 7 days\n"
             "Captures id, name, item_no, created_at."),
            ("Step 2 - Sales aggregates (sales CTE)",
             "today_qty = SUM(quantity) where order created_at >= now - 1 day\n"
             "qty_7d    = SUM(quantity) where order created_at >= now - 7 days"),
            ("Step 3 - Velocity ratio per item",
             "velocity_ratio = today_qty / (qty_7d / 7)\n"
             "'How many times the recent daily average is today selling at?'"),
            ("Step 4 - Filter",
             "qty_7d > 0\n"
             "AND velocity_ratio > 2.0"),
            ("Step 5 - Age + rank",
             "age_days = (NOW - created_at) / 86400\n"
             "Order by velocity_ratio DESC, return top PRODUCT_LIMIT.\n"
             "Launch age is humanized in Python (e.g. '3 days old')."),
        ],
        "example_note": "Today = Day 7. Threshold = velocity_ratio > 2.0.",
        "example_headers": ["item", "created (days ago)", "qty_7d", "today_qty", "daily_avg (qty_7d/7)", "velocity_ratio", "> 2.0?"],
        "example_rows": [
            ["A", "3", "14", "10", "2.0", "5.0", "QUALIFIES"],
            ["B", "5", "21", "3", "3.0", "1.0", "no (today < average)"],
            ["C", "2", "0", "4", "n/a", "filtered", "no (qty_7d = 0)"],
        ],
        "example_concl": "Item A is launching with momentum: lifetime daily pace ~2/day, but today did 10 -> 5x pace. Surface it.",
        "outputs": [
            ("today_sales", "quantity sold in the last 24h"),
            ("sales_7d", "quantity sold in the last 7 days"),
            ("velocity_ratio", "today_qty / (qty_7d / 7)"),
            ("launched", "humanised age, e.g. '3 days old'"),
            ("score", "the velocity_ratio value (used for ranking)"),
        ],
        "note": (
            "The 2.0 velocity threshold and the 7-day age window are hard-coded in "
            "queries.py. The today / 7d / 90d windows are scaled by the WINDOW_MULT env var "
            "(default 1)."
        ),
    },
]

# ============================================================================
# AI CONTEXT (verbatim from RULE_AI_CONTEXT in app/rules/ai.py)
# Each rule's static "AI Detection Logic" (the AI Insight) + "AI Recommendation
# Reason". These are the grounding strings fed to Gemini; when AI is disabled or
# fails, this exact text is returned to the merchant verbatim.
# ============================================================================
AI_CONTEXT = {
    "viral": {
        "detection": "Click-through rate is above average by 50%, but total traffic is low.",
        "reason": "High CTR proves strong appeal; the product only needs more traffic exposure to become a bestseller.",
        "example": (
            "Product A has a Click-Through Rate (CTR) of 50%, about 2.5 times the store "
            "average, yet only 10 Page Views (PV). Increase traffic with ads or homepage "
            "placement to convert this strong appeal into sales and lift Conversion Rate "
            "(CR) by ~15%."
        ),
    },
    "clearance": {
        "detection": "Inventory exceeds 100 units while 30-day sales are below 5.",
        "reason": "Inventory accumulation is severe; promotional campaigns are recommended for faster liquidation.",
        "example": (
            "Product A (500 units, only 3 sold in 30 days) and Product B (800 units, ~1,200 "
            "Days of Inventory) are tying up stock. Launch a clearance promotion or bundle "
            "to liquidate them and improve Inventory Turnover by ~25%."
        ),
    },
    "hidden_gem": {
        "detection": "Impressions are low, but conversions occur whenever visitors arrive.",
        "reason": "Extremely strong conversion ability but low visibility; this is an undiscovered opportunity product.",
        "example": (
            "Product A converts at 16.7%, about 4.2 times the store average Conversion Rate "
            "(CR), but had only 12 impressions. Give it more ad spend and visibility to "
            "unlock an estimated Return on Ad Spend (ROAS) gain of ~30%."
        ),
    },
    "high_attention": {
        "detection": "Product page traffic ranks in the store's top 5%, but add-to-cart rate is below average.",
        "reason": "Many users view but do not buy, usually due to pricing or decision barriers; discounts or incentives may help conversion.",
        "example": (
            "Products 1033468 and 1033460 are among the most-viewed pages (25 and 21 Page "
            "Views) yet have a 0% Add-to-Cart Rate (ATC). Try a discount, clearer pricing, "
            "or better images to lift the Add-to-Cart Rate (ATC) by ~20%."
        ),
    },
    "basket_magnet": {
        "detection": "Frequently appears together with other products in the same order.",
        "reason": "This product naturally drives bundle purchases and is ideal for increasing average order value.",
        "example": (
            "Product A appears in a basket with other items in 80% of its 50 orders. Feature "
            "it in 'frequently bought together' and bundle offers to raise Average Order "
            "Value (AOV) by ~18%."
        ),
    },
    "loyalty": {
        "detection": "Products with the highest repeat purchase counts among existing customers.",
        "reason": "Product quality and customer satisfaction are excellent, reducing customer acquisition costs through word-of-mouth.",
        "example": (
            "Product A has a 35% repeat-purchase rate (35 of 100 buyers came back). Promote "
            "it through subscription or replenishment offers to grow customer Lifetime Value "
            "(LTV) by ~22%."
        ),
    },
    "profit": {
        "detection": "Gross margin exceeds 50%.",
        "reason": "Products with the highest profit potential are suitable for higher advertising budgets.",
        "example": (
            "Product A keeps a 70% gross margin (sells at 100, costs 30). It can absorb a "
            "higher advertising budget while still protecting profit, supporting a Net "
            "Profit gain of ~12%."
        ),
    },
    "engagement": {
        "detection": "Average user browsing time on the product page exceeds 90 seconds.",
        "reason": "Customers deeply engage with the page, indicating strong content attraction and exploration intent.",
        "example": (
            "Product A holds visitors for an average of 145 seconds across 12 customers, well "
            "above the 90-second bar. Add related products or a strong call-to-action to "
            "convert this attention and extend Dwell Time by ~35%."
        ),
    },
    "momentum": {
        "detection": "Sales growth curve is steepest within the first 7 days after launch.",
        "reason": "Newly launched product is rapidly gaining traction and market attention.",
        "example": (
            "Water Bottle (velocity ratio 5.96), City Wallpapers (2.45), and Football (2.21) "
            "are rapidly gaining traction. Promote these new products to new customers to "
            "boost new customer acquisition by ~20%."
        ),
    },
}

# ============================================================================
# RENDER EACH RULE
# ============================================================================
for idx, rule in enumerate(RULES):
    h1("Rule {} - {}".format(rule["id"], rule["title"]))

    make_table(
        ["Field", "Value"],
        [
            ["Rule ID", rule["id"]],
            ["AI Recommendation Tag", rule["title"]],
            ["Internal Tag", rule["tag"]],
            ["KPI Target", rule["kpi"]],
        ],
    )

    h2("What this rule detects")
    para(rule["detects"])

    h2("Rule definition (trigger formula)")
    code_block(rule["formula"])

    h2("Step-by-step calculation")
    for step_title, step_body in rule["steps"]:
        h3(step_title)
        code_block(step_body)

    h2("Worked example - why an item appears")
    para(rule["example_note"])
    make_table(rule["example_headers"], rule["example_rows"])
    para("Result: " + rule["example_concl"], italic=True)

    h2("Output fields (per product)")
    make_table(["Field", "Meaning"], [[f, m] for f, m in rule["outputs"]])

    # ---- AI INSIGHT — the live, AI-written recommendation (from ai.py enrich_reason) ----
    aictx = AI_CONTEXT.get(rule["tag"])
    if aictx and aictx.get("example"):
        h2("AI Insight")
        p = para()
        r = p.add_run("AI INSIGHT")
        r.bold = True
        r.font.size = Pt(10)
        r.font.color.rgb = RGBColor(0xC0, 0x5A, 0x2B)  # warm orange, like the UI box
        insight_box(aictx["example"])

    if rule.get("note"):
        h2("Notes")
        para(rule["note"])

    if idx != len(RULES) - 1:
        doc.add_page_break()

# ============================================================================
# CLOSING - BENEFITS
# ============================================================================
doc.add_page_break()
h1("Benefits")
para("This architecture:")
for b in [
    "Keeps token usage low",
    "Reduces AI hallucinations",
    "Improves recommendation accuracy",
    "Lowers operational cost",
    "Provides explainable and auditable results",
    "Maintains fast response times",
    "Scales efficiently to 100,000+ products",
    "Ensures business logic remains deterministic and controllable",
]:
    bullet(b)
para(
    "In summary, the database determines which products deserve attention, while AI "
    "determines which of the top products should be recommended and explains why."
)

OUT = r"d:\B2C\_B2c_php\works\b2c-ai-assistant\docs\AI_Product_Recommendation_Rules.docx"
doc.save(OUT)
print("Saved:", OUT)
